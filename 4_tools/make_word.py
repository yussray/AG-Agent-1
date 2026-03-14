# ============================================================
# make_word.py — Antigravity Multi-Agent System
# Generate Word (.docx) documents
# ============================================================

import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
import paths  # noqa: E402

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from config import OUTPUT_DIR


def make_word(content: str, title: str = "Document", filename: str = None) -> str:
    """
    Generate a .docx file from markdown-like text.

    Args:
        content: The document text (## headings, - bullets supported)
        title: Document title
        filename: Output filename

    Returns:
        Absolute path to saved .docx file
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    if not filename:
        slug = re.sub(r'[^\w]', '_', title.lower())[:25]
        filename = f"doc_{slug}.docx"

    path = os.path.join(OUTPUT_DIR, filename)
    doc = Document()

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for line in content.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)

    doc.save(path)
    return os.path.abspath(path)
